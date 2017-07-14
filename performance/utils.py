# -*- coding: UTF-8 -*-

import os
import docx
import xlrd
from django.conf import settings
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator


class Utils(object):
    """工具类"""

    @staticmethod
    def get_contents(file_path):
        if not os.path.exists(file_path):
            raise Exception("File not exist: {0}".format(file_path))

        content = ""
        extension = Utils.file_extension(file_path)
        if extension in settings.EXCEL_EXTENSIONS:
            content = Utils.excel_get_contents(file_path)
        elif extension in settings.DOCX_EXTENSIONS:
            content = Utils.docx_get_contents(file_path)
        elif extension in settings.PDF_EXTENSIONS:
            content = Utils.pdf_get_contents(file_path)
        elif extension in [".txt", ".csv", ".html", ".xml", ".json"]:
            content = Utils.file_get_contents(file_path)

        return content

    @staticmethod
    def file_get_contents(file_path):
        content = ""
        with open(file_path, 'r') as f:
            content += f.read()

        return content

    @staticmethod
    def excel_get_contents(file_path):
        content = ""
        excel = xlrd.open_workbook(file_path)
        for sheet in excel.sheets():
            for i in range(sheet.nrows):
                content += ','.join(
                    [str(item) for item in sheet.row_values(i)])

        return content

    @staticmethod
    def docx_get_contents(file_path):
        content = ""
        doc = docx.Document(file_path)
        paras = doc.paragraphs
        for p in paras:
            content += p.text

        return content

    @staticmethod
    def pdf_get_contents(file_path):
        try:
            fp = open(file_path, "rb")  # 获取文档对象，你把algorithm.pdf换成你自己的文件名即可。
            parser = PDFParser(fp)  # 创建一个与文档相关联的解释器
            doc = PDFDocument(parser)  # PDF文档对象
            parser.set_document(doc)  # 链接解释器和文档对象
            resource = PDFResourceManager()  # 创建PDF资源管理器
            laparam = LAParams()  # 参数分析器
            device = PDFPageAggregator(resource, laparams=laparam)  # 创建一个聚合器
            interpreter = PDFPageInterpreter(resource, device)  # 创建PDF页面解释器
            content = ""
            for page in PDFPage.create_pages(doc):  # 使用文档对象得到页面集合
                interpreter.process_page(page)  # 使用页面解释器来读取
                layout = device.get_result()  # 使用聚合器来获取内容
                for out in layout:
                    if hasattr(out, "get_text"):
                        content += out.get_text()
        except Exception as e:
            raise e
        finally:
            if fp:
                fp.close()

        return content

    @staticmethod
    def file_extension(file_path):
        return os.path.splitext(file_path)[1].lower()

    @staticmethod
    def file_put_contents(file_path, contents):
        with open(file_path, 'w') as f:
            f.write(contents)

        return file_path
